import os
from pydoc import doc
import re
import json
import argparse
#from tkinter import font
from turtle import title
from typing import List, Dict
from datetime import datetime

# Environment & AI API
from dotenv import load_dotenv
from openai import OpenAI
import openai
import requests

# File handling
import pandas as pd
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Optional PDF support
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
except ImportError:
    pdf_extract_text = None

# Optional NLP
try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
except Exception:
    nlp = None

load_dotenv()

# ---------------------- CONFIG ----------------------

BRD_SECTIONS = [
    "Executive Summary",
    "Project Overview",
    "Scope (In / Out)",
    "Business Objectives",
    "Functional Requirements",
    "Non-Functional Requirements",
    "Acceptance Criteria",
    "Assumptions & Constraints",
    "Risks",
    "Stakeholders",
    "Data Flow & Integration",
    "Milestones & Timeline",
    "Change Control",
    "Glossary",
    "Appendices",
]

# ---------------------- FILE PARSER ----------------------


class FileProcessor:
    """Handles extraction of text from various file formats."""

    @staticmethod
    def extract_text(file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()

        if ext in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        elif ext == ".docx":
            doc = DocxDocument(file_path)
            return "\n".join([p.text for p in doc.paragraphs])

        elif ext in [".csv", ".xlsx", ".xls"]:
            df = pd.read_csv(file_path) if ext == ".csv" else pd.read_excel(file_path)
            return df.to_string(index=False)

        elif ext == ".json":
            with open(file_path, "r", encoding="utf-8") as f:
                return json.dumps(json.load(f), indent=2)

        elif ext == ".pdf" and pdf_extract_text:
            return pdf_extract_text(file_path)

        else:
            raise ValueError(f"Unsupported file type or missing parser for {ext}")


# ---------------------- FACT EXTRACTION ----------------------


class FactExtractor:
    """Extracts key facts like dates, stakeholders, risks, and features."""

    @staticmethod
    def extract_keywords(text: str) -> Dict[str, List[str]]:
        facts = {
            "projects": [],
            "dates": [],
            "standards": [],
            "stakeholders": [],
            "risks": [],
            "features": [],
            "metrics": [],
            "systems": [],
        }

        # Projects
        proj_match = re.findall(
            r"\b[A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)*\s+(?:System|Project)\b", text
        )
        facts["projects"].extend(proj_match)

        # Dates
        date_match = re.findall(
            r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},\s+\d{4}"
            r"|\b\d{4}-\d{2}-\d{2}\b"
            r"|\b\d{1,2}/\d{1,2}/\d{4}\b",
            text,
        )
        facts["dates"].extend(date_match)

        # Standards
        iso_match = re.findall(r"\bISO\s*\d{4,6}\b", text)
        facts["standards"].extend(iso_match)

        # Stakeholders
        stake_match = re.findall(
            r"(Manager|Lead|Engineer|Director|Sponsor)", text, re.IGNORECASE
        )
        facts["stakeholders"].extend(set(stake_match))

        # Risks
        if "risk" in text.lower():
            risk_lines = [
                line.strip() for line in text.splitlines() if "risk" in line.lower()
            ]
            facts["risks"].extend(risk_lines)

        # Features
        feat_lines = [
            line.strip()
            for line in text.splitlines()
            if any(kw in line.lower() for kw in ["feature", "shall", "must", "requirement"])
        ]
        facts["features"].extend(feat_lines)

        # Systems
        sys_match = re.findall(r"\bSAP\b|\bMES\b|\bERP\b", text)
        facts["systems"].extend(sys_match)

        # Metrics
        perc_match = re.findall(r"\b\d{1,3}%\b", text)
        facts["metrics"].extend(perc_match)

        # Deduplicate
        for k in facts:
            facts[k] = list(set(facts[k]))

        return facts


# ---------------------- RAG CHUNKING ----------------------


def chunk_text(text: str, max_chars=1500) -> List[str]:
    """Splits large text into manageable chunks for AI processing."""
    paras = text.split("\n")
    chunks, current = [], ""

    for para in paras:
        if len(current) + len(para) + 1 <= max_chars:
            current += para + "\n"
        else:
            chunks.append(current.strip())
            current = para + "\n"

    if current.strip():
        chunks.append(current.strip())

    return chunks


def find_relevant_chunks(section: str, chunks: List[str], facts: Dict[str, List[str]], top_k=5) -> List[str]:
    """Find chunks most relevant to a given BRD section using keyword overlap."""
    scored = []
    keywords = section.lower().split() + sum(facts.values(), [])
    for chunk in chunks:
        score = sum(1 for kw in keywords if kw.lower() in chunk.lower())
        scored.append((score, chunk))

    scored.sort(key=lambda x: x[0], reverse=True)
    return [chunk for score, chunk in scored[:top_k] if score > 0] or chunks[:1]


# ---------------------- AI PROCESSOR ----------------------


class AIProcessor:
    """Handles interaction with AI models (OpenAI, Ollama, etc.)."""

    def __init__(self):
        self.api_type = os.getenv("API_TYPE", "openai")
        self.openai_model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        self.temperature = float(os.getenv("AI_TEMPERATURE", "0.2"))

        if self.api_type == "openai":
            self.client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

    def call_ai(self, system_prompt: str, user_prompt: str) -> str:
        if self.api_type == "none" or os.getenv("DRY_RUN_NO_API") == "1":
            return "[Generated content based on extracted facts only]"

        if self.api_type == "openai":
            openai.api_key = os.getenv("OPENAI_API_KEY")
            resp = self.client.chat.completions.create(
                model=self.openai_model,
                temperature=self.temperature,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
            )
            return resp.choices[0].message.content

        elif self.api_type == "ollama":
            base_url = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
            model = os.getenv("OLLAMA_MODEL", "llama3")
            resp = requests.post(
                f"{base_url}/api/generate",
                json={"model": model, "prompt": f"{system_prompt}\n\n{user_prompt}"},
            )
            return resp.json().get("response", "")

        else:
            raise ValueError(f"Unsupported API_TYPE {self.api_type}")


# ---------------------- BRD GENERATOR ----------------------


class BRDGenerator:
    """Generates a BRD document from extracted text and AI responses."""

    def __init__(self, ai: AIProcessor):
        self.ai = ai

    def generate_brd(self, text: str) -> Dict[str, str]:
        facts = FactExtractor.extract_keywords(text)
        chunks = chunk_text(text)

        brd = {}
        for section in BRD_SECTIONS:
            relevant = find_relevant_chunks(section, chunks, facts, top_k=5)
            system_prompt = (
                '''You are a senior business analyst.
Generate a clear, concise section for a Business Requirements Document (BRD) 
using ONLY the provided source content and extracted facts. 

Guidelines:
- Use the exact provided project title: "{project_title}" (do not invent or change it).
- Do NOT repeat the section name inside the response; begin directly with the content.
- Avoid duplicating ideas already covered in other sections unless strictly necessary.
- Write in clear, professional business English.
- Keep the overall BRD length under 20 pages; summarize and be concise.
- If information is missing, state 'Not specified in provided input'.
- ABSOLUTELY DO NOT include or repeat the section title in your response. Start IMMEDIATELY with the body content.'''
            )
            user_prompt = (
                 f"Write the section titled '{section}' based on the following inputs:\n\n"
                 f"Source Chunks:\n{json.dumps(relevant, indent=2)}\n\n"
                 f"Extracted Facts:\n{json.dumps(facts, indent=2)}\n\n"
                 "Follow the system prompt instructions strictly."
            )
            content = self.ai.call_ai(system_prompt, user_prompt)
            # Strip leading section title if present
            title_lower = section.lower()
            if content.lower().startswith(title_lower):
                content = content[len(section):].lstrip()  # Remove the repeating title
            brd[section] = self._clean_markdown_artifacts(content)

        return brd
    
    def review_and_compress(self, brd: Dict[str, str], project_title: str) -> Dict[str, str]:
        combined = "\n\n".join(f"{sec}:\n{txt}" for sec, txt in brd.items())
        review_prompt = f"""
    You are a senior business analyst reviewing a full Business Requirements Document (BRD).

    Tasks:
        1. Remove repeated subheadings or duplicate content across sections.
        2. Ensure document length is concise (under 20 pages).
        3. Keep project title consistent as "{project_title}".
        4. Do not remove required sections, but compress verbose explanations.
        5. Maintain professional tone.

    Here is the draft BRD:
    {combined}
    Return the reviewed BRD as JSON with section names as keys and content as values, e.g., {{"Executive Summary": "content...", "Project Overview": "content..."}}.
    """
        reviewed_response = self.ai.call_ai("You are a document reviewer.", review_prompt)
        try:
            reviewed_brd = json.loads(reviewed_response)
        except json.JSONDecodeError:
            reviewed_brd = brd  # Fallback to original if parsing fails
        return reviewed_brd
    
    def _clean_markdown_artifacts(self, text: str) -> str:
        if not text:
            return "Not specified in provided input."

        text = re.sub(r"\*\*\*(.*?)\*\*\*", r"\1", text)
        text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
        text = re.sub(r"\*(.*?)\*", r"\1", text)
        text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
        text = re.sub(r"^>\s*", "", text, flags=re.MULTILINE)
        text = re.sub(r"^\s*[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()

    def save_brd_docx(self, brd: Dict[str, str], output_path: str):
        doc = DocxDocument()

    # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

    # Title Page
        title = doc.add_heading("Business Requirements Document", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(0, 0, 255)  # Blue color for title
        doc.add_paragraph("", style="Normal")

        subtitle = doc.add_paragraph("Generated by AI-Powered BRD Generator\n")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].italic = True

        date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Table of Contents
        doc.add_heading("Table of Contents", level=1)
        doc.paragraphs[-1].runs[0].font.color.rgb = RGBColor(0, 0, 255)  # Blue for TOC heading

        page_num = 1
        for section_title in brd.keys():
            doc.add_paragraph(f"{section_title} ........ {page_num}")
            page_num += 1

        doc.add_page_break()

        # BRD Sections
        for section_title, content in brd.items():
            clean_content = self._clean_markdown_artifacts(content)
            heading = doc.add_heading(section_title, level=1)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 255)  # Blue color for section headers

            for para_text in clean_content.split("\n"):
                para_text = para_text.strip()
                if para_text:
                    p = doc.add_paragraph(para_text, style="Normal")
                    p.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    p.space_after = Pt(6)
            doc.add_paragraph("")

        # Safe Save
        try:
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)

            doc.save(output_path)
            print(f"✅ BRD successfully saved to {output_path}")

        except PermissionError:
            print(f"❌ Permission denied: Close '{output_path}' in Word and try again.")
        except Exception as e:
            print(f"❌ Failed to save document: {str(e)}")


# ---------------------- CLI ----------------------

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-file", required=True)
    parser.add_argument("--output-file", default="BRD.docx")
    args = parser.parse_args()

    raw_text = FileProcessor.extract_text(args.input_file)
    ai = AIProcessor()
    gen = BRDGenerator(ai)
    brd = gen.generate_brd(raw_text)
    gen.save_brd_docx(brd, args.output_file)

    print(f"BRD saved to {args.output_file}")
