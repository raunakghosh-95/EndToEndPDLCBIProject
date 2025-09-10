import streamlit as st
import os
import tempfile
import io
import json
import re
from datetime import datetime
from dotenv import load_dotenv
from brd_generator_rag import FileProcessor, AIProcessor, BRDGenerator
from system_design_agent import SystemDesignAgent, DiagramRenderer
from docx import Document as DocxDocument

# NEW: imports for Power BI tab
import uuid
from urllib.parse import urlparse, parse_qs, urlencode, urlunparse
import streamlit.components.v1 as components
from typing import Dict
# NEW: imports for rendering diagrams
import base64
import zlib
import requests
import streamlit_mermaid as stmd

load_dotenv()

# Page configuration
st.set_page_config(
    page_title="BRD Generator",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        text-align: center;
        padding: 2rem 0;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    .upload-section {
        border: 2px dashed #cccccc;
        padding: 2rem;
        border-radius: 10px;
        text-align: center;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        color: #155724;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    .error-box {
        background-color: #f8d7da;
        border: 1px solid #f5c6cb;
        color: #721c24;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    /* Make the Power BI iframe responsive */
    .pbicontainer {
        position: relative;
        width: 90%;
        padding-bottom: 62.25%; /* 16:9 -> 56.25%, tweaked for Power BI default 600x373.5 */
        height: 0;
    }
    .pbicontainer iframe {
        position: absolute;
        top: 0; left: 0;
        width: 90%;
        height: 100%;
        border: 0;
    }
</style>
""", unsafe_allow_html=True)


def initialize_session_state():
    """Initialize session state variables"""
    if 'brd_generated' not in st.session_state:
        st.session_state.brd_generated = False
    if 'brd_content' not in st.session_state:
        st.session_state.brd_content = None
    if 'brd_filename' not in st.session_state:
        st.session_state.brd_filename = None
    if 'custom_template' not in st.session_state:
        st.session_state.custom_template = None
    if 'template_sections' not in st.session_state:
        st.session_state.template_sections = None
    # NEW: defaults for Power BI tab controls
    if 'pbi_autorefresh_enabled' not in st.session_state:
        st.session_state.pbi_autorefresh_enabled = False
    if 'pbi_autorefresh_seconds' not in st.session_state:
        st.session_state.pbi_autorefresh_seconds = 120
    if 'pbi_height_px' not in st.session_state:
        st.session_state.pbi_height_px = 650
    if 'pbi_width_px' not in st.session_state:
        st.session_state.pbi_width_px = 900
    if 'pbi_cache_bust' not in st.session_state:
        st.session_state.pbi_cache_bust = True
    # NEW: for system design
    if 'system_design_generated' not in st.session_state:
        st.session_state.system_design_generated = False
    if 'system_design_artifacts' not in st.session_state:
        st.session_state.system_design_artifacts = None
    # NEW: for dashboard dev doc
    if 'dashboard_dev_generated' not in st.session_state:
        st.session_state.dashboard_dev_generated = False
    if 'dashboard_dev_content' not in st.session_state:
        st.session_state.dashboard_dev_content = None
    if 'dashboard_dev_filename' not in st.session_state:
        st.session_state.dashboard_dev_filename = None


def validate_environment():
    """Check if required environment variables are set"""
    required_vars = ['OPENAI_API_KEY']
    missing_vars = []

    for var in required_vars:
        if not os.getenv(var):
            missing_vars.append(var)

    return missing_vars


def extract_template_sections(template_path):
    """Extract section headings from a BRD template document"""
    try:
        if template_path.endswith('.docx'):
            doc = DocxDocument(template_path)
            sections = []

            # Extract headings from DOCX
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    text = paragraph.text.strip()
                    if text and len(text) > 0:
                        sections.append(text)

            return sections

        elif template_path.endswith('.txt') or template_path.endswith('.md'):
            with open(template_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Extract markdown-style headings or numbered sections
            sections = []
            lines = content.split('\n')

            for line in lines:
                line = line.strip()
                # Match markdown headings (# ## ###) or numbered sections (1. 2. etc.)
                if (line.startswith('#') or
                    re.match(r'^\d+\.?\s+[A-Z]', line) or
                    line.isupper() and len(line.split()) <= 5):
                    clean_section = re.sub(r'^#+\s*|\d+\.?\s*', '', line).strip()
                    if clean_section:
                        sections.append(clean_section)

            return sections

        elif template_path.endswith('.json'):
            with open(template_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            # Assume JSON has sections as keys or in a sections array
            if isinstance(data, dict):
                if 'sections' in data:
                    return data['sections']
                else:
                    return list(data.keys())
            elif isinstance(data, list):
                return data

        return None

    except Exception as e:
        st.error(f"Error extracting template sections: {str(e)}")
        return None


class TemplateBRDGenerator(BRDGenerator):
    """Extended BRD Generator that supports custom templates"""

    def __init__(self, ai: AIProcessor, custom_sections=None):
        super().__init__(ai)
        self.custom_sections = custom_sections

    def get_sections(self):
        """Get sections to use - custom template or default"""
        if self.custom_sections:
            return self.custom_sections
        else:
            # Default sections from original code
            return [
                "Executive Summary",
                "Project Overview",
                "Scope (In / Out)",
                "Business Objectives",
                "Functional Requirements",
                "Non-Functional Requirements",
                "Acceptance Criteria",
                "Assumptions & Constraints",
                "Risks",
                "Stakeholders",
                "Data Flow & Integration",
                "Milestones & Timeline",
                "Change Control",
                "Glossary",
                "Appendices"
            ]
            
    def review_and_compress(self, brd: dict, project_title: str) -> str:
        """Run a review pass to deduplicate, compress, and enforce page limits."""
        combined = "\n\n".join(f"{sec}:\n{txt}" for sec, txt in brd.items())
        review_prompt = f"""
You are a senior business analyst reviewing a full Business Requirements Document (BRD).

Tasks:
    1. Scan for and remove ANY duplicated content, phrases, or ideas across sections (e.g., avoid repeating data flow details in Overview and Scope).
    2. Ensure NO section exceeds 300 words; aggressively summarize while retaining key facts.
    3. Enforce total document length under 20 pages (~5000 words max) by condensing verbose text.
    4. Keep project title consistent as "{project_title}".
    5. Do not remove required sections, but merge similar ideas.
    6. Maintain professional tone and eliminate redundancy.

Here is the draft BRD:
{combined}
"""
        reviewed = self.ai.call_ai("You are a document reviewer.", review_prompt)
        word_count = len(reviewed.split())
        if word_count > 5000:  # ~20 pages at 250 words/page
            review_prompt += f"\n\nCurrent word count: {word_count}. Reduce by at least 30% without losing key facts."
            reviewed = self.ai.call_ai("You are a document reviewer.", review_prompt)
        return reviewed

    def generate_brd(self, text: str) -> dict:
        """Generate BRD using custom or default sections"""
        from brd_generator_rag import FactExtractor, chunk_text, find_relevant_chunks

        facts = FactExtractor.extract_keywords(text)
        chunks = chunk_text(text)

        sections = self.get_sections()
        brd = {}

        for section in sections:
            relevant = find_relevant_chunks(section, chunks, facts, top_k=5)
            system_prompt = (
                "You are a senior business analyst. "
                "Generate a detailed section for a Business Requirements Document (BRD) "
                "using ONLY the provided source content and extracted facts. "
                "If information is missing, state 'Not specified in provided input'. "
                "Write in clear professional business English. "
                f"Focus specifically on the '{section}' section requirements."
            )
            user_prompt = (
                f"Section: {section}\n\n"
                f"Source Chunks:\n{json.dumps(relevant, indent=2)}\n\n"
                f"Extracted Facts:\n{json.dumps(facts, indent=2)}"
            )
            brd[section] = self.ai.call_ai(system_prompt, user_prompt)

        return brd


def save_uploaded_file(uploaded_file):
    """Save uploaded file to temporary location"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            return tmp_file.name
    except Exception as e:
        st.error(f"Error saving uploaded file: {str(e)}")
        return None


def generate_brd_from_file(file_path, progress_callback=None, template_sections=None):
    """Generate BRD from uploaded file with optional template sections"""
    try:
        if progress_callback:
            progress_callback(20, "Extracting text from file...")

        # Extract text from file
        raw_text = FileProcessor.extract_text(file_path)

        if not raw_text.strip():
            raise ValueError("No text could be extracted from the uploaded file")

        if progress_callback:
            progress_callback(40, "Initializing AI processor...")

        # Initialize AI processor with custom template
        ai = AIProcessor()
        gen = TemplateBRDGenerator(ai, template_sections)

        if progress_callback:
            progress_callback(60, "Generating BRD sections...")

        # Generate BRD
        brd = gen.generate_brd(raw_text)

        if progress_callback:
            progress_callback(80, "Creating document...")

        # Create temporary file for BRD
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            gen.save_brd_docx(brd, tmp_file.name)

            if progress_callback:
                progress_callback(100, "BRD generated successfully!")

            return tmp_file.name, brd

    except Exception as e:
        raise Exception(f"Error generating BRD: {str(e)}")


def system_design_interface():
    """System Design Generator Interface"""

    st.header("🏗️ System Design Generator")
    st.markdown("Generate system architecture and UML diagrams from your BRD content using AI")

    # Input methods
    input_method = st.radio(
        "Choose input method:",
        ["Use Generated BRD", "Paste BRD Text", "Upload BRD File"],
        horizontal=True
    )

    brd_text = ""

    if input_method == "Use Generated BRD":
        if st.session_state.brd_generated and st.session_state.brd_content:
            # Convert BRD dict to text
            brd_text = ""
            for section, content in st.session_state.brd_content.items():
                brd_text += f"{section}:\n{content}\n\n"

            with st.expander("📖 BRD Content to be Used"):
                st.text_area("BRD Content", value=brd_text[:1000] + "..." if len(brd_text) > 1000 else brd_text, height=200, disabled=True)
        else:
            st.warning("No BRD has been generated yet. Please generate a BRD first or use another input method.")
            return

    elif input_method == "Paste BRD Text":
        brd_text = st.text_area(
            "Paste your BRD content here:",
            height=300,
            placeholder="Executive Summary: ...\n\nFunctional Requirements: ..."
        )

    elif input_method == "Upload BRD File":
        uploaded_brd = st.file_uploader(
            "Upload BRD file",
            type=['txt', 'md', 'docx'],
            help="Upload a text file containing your BRD"
        )

        if uploaded_brd:
            try:
                if uploaded_brd.type == "text/plain":
                    brd_text = str(uploaded_brd.read(), "utf-8")
                else:
                    # Save and process file
                    temp_path = save_uploaded_file(uploaded_brd)
                    if temp_path:
                        from brd_generator_rag import FileProcessor
                        brd_text = FileProcessor.extract_text(temp_path)
                        os.unlink(temp_path)
            except Exception as e:
                st.error(f"Error reading file: {str(e)}")

    # Generate button
    if st.button("🚀 Generate System Design", type="primary", disabled=not brd_text.strip()):
        if not brd_text.strip():
            st.error("Please provide BRD content first")
            return

        try:
            # Initialize progress tracking
            progress_bar = st.progress(0)
            status_text = st.empty()

            # Initialize the AI agent
            status_text.text("🤖 Initializing System Design Agent...")
            progress_bar.progress(10)

            agent = SystemDesignAgent()

            # Generate all artifacts
            status_text.text("🔍 Analyzing BRD content...")
            progress_bar.progress(20)

            with st.spinner("Generating system design artifacts..."):
                artifacts = agent.generate_all_artifacts(brd_text)

            progress_bar.progress(100)
            status_text.text("✅ System design artifacts generated!")

            # Store in session state
            st.session_state.system_design_generated = True
            st.session_state.system_design_artifacts = artifacts

            st.success("🎉 System design artifacts generated successfully!")

        except Exception as e:
            st.error(f"Error generating system design: {str(e)}")
            st.info("Please check your API configuration and BRD content.")


def display_system_design_results():
    """Display generated system design artifacts"""

    if not hasattr(st.session_state, 'system_design_artifacts'):
        return

    artifacts = st.session_state.system_design_artifacts

    st.divider()
    st.header("📊 Generated System Design Artifacts")

    # Analysis Summary
    with st.expander("🔍 BRD Analysis Summary", expanded=False):
        analysis = artifacts.get('analysis', {})
        if isinstance(analysis, dict):
            st.json(analysis)
        else:
            st.text(str(analysis))

    # Diagram tabs
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "🏗️ System Architecture",
        "👤 Use Cases",
        "⚡ Sequence",
        "📦 Classes",
        "🗃️ Entity Relationship",
        "📚 Documentation"
    ])

    with tab1:
        st.subheader("System Architecture Diagram")
        system_arch = artifacts.get('system_architecture', '')
        if system_arch:
            # Clean the Mermaid code to remove ```mermaid and ``` markers
            cleaned_system_arch = '\n'.join(line for line in system_arch.split('\n') if not line.strip().startswith('```') and line.strip())
            # Display Mermaid diagram natively
            try:
                stmd.st_mermaid(cleaned_system_arch, height=600, width=800)
            except Exception as e:
                st.error(f"Error rendering Mermaid diagram: {str(e)}")
                st.code(system_arch, language="text")

            # Copy-paste section
            with st.expander("📋 Copy Mermaid Code"):
                st.code(system_arch, language="text")
                st.info("💡 Copy this code to [Mermaid Live Editor](https://mermaid.live) to render and export")

    with tab2:
        st.subheader("Use Case Diagram")
        use_case = artifacts.get('use_case_diagram', '')
        if use_case:
            st.code(use_case, language="text")
            st.info("💡 Copy this PlantUML code to [PlantUML Online](https://www.plantuml.com/plantuml) to render")

    with tab3:
        st.subheader("Sequence Diagram")
        sequence = artifacts.get('sequence_diagram', '')
        if sequence:
            st.code(sequence, language="text")
            st.info("💡 Copy this PlantUML code to [PlantUML Online](https://www.plantuml.com/plantuml) to render")

    with tab4:
        st.subheader("Class Diagram")
        class_diagram = artifacts.get('class_diagram', '')
        if class_diagram:
            st.code(class_diagram, language="text")
            st.info("💡 Copy this PlantUML code to [PlantUML Online](https://www.plantuml.com/plantuml) to render")

    with tab5:
        st.subheader("Entity Relationship Diagram")
        er_diagram = artifacts.get('er_diagram', '')
        if er_diagram:
    # Try to match ```mermaid ... ```
            match = re.search(r"```mermaid\s*([\s\S]*?)(```|$)", er_diagram)
            if match:
                cleaned_er_diagram = match.group(1).strip()
            else:
        # fallback: remove just the word 'mermaid' if present
                cleaned_er_diagram = er_diagram.replace("```mermaid", "").replace("```", "").strip()
            try:
                stmd.st_mermaid(cleaned_er_diagram, height=600, width=800)
            except Exception as e:
                st.error(f"Error rendering Mermaid ER diagram: {str(e)}")
                st.code(er_diagram, language="text")

            with st.expander("📋 Copy Mermaid ER Code"):
                st.code(er_diagram, language="text")
                st.info("💡 Copy this code to [Mermaid Live Editor](https://mermaid.live) to render and export")

    with tab6:
        st.subheader("Data Model Documentation")
        data_docs = artifacts.get('data_model_docs', {})
        if isinstance(data_docs, dict):
            st.json(data_docs)
        else:
            st.text(str(data_docs))

    # Download section
    st.divider()
    col1, col2, col3 = st.columns(3)

    with col1:
        # Download JSON artifacts
        artifacts_json = json.dumps(artifacts, indent=2)
        st.download_button(
            label="📥 Download Artifacts (JSON)",
            data=artifacts_json,
            file_name=f"system_design_artifacts_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            mime="application/json",
            type="secondary",
            use_container_width=True
        )

    with col2:
        # Download HTML preview
        try:
            renderer = DiagramRenderer()
            html_preview = renderer.generate_html_preview(artifacts)
            st.download_button(
                label="📥 Download HTML Preview",
                data=html_preview,
                file_name=f"system_design_preview_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
                mime="text/html",
                type="secondary",
                use_container_width=True
            )
        except Exception as e:
            st.error(f"Error generating HTML preview: {str(e)}")

    with col3:
        # Reset button
        if st.button("🔄 Generate New Design", type="secondary", use_container_width=True):
            if hasattr(st.session_state, 'system_design_generated'):
                del st.session_state.system_design_generated
            if hasattr(st.session_state, 'system_design_artifacts'):
                del st.session_state.system_design_artifacts
            st.rerun()


# NEW: helper to add a cache-busting param to the Power BI report URL
def _cache_busted_url(src_url: str) -> str:
    """
    Append a short-lived cache-busting query parameter (cb) to the URL.
    This forces the iframe to reload and pick up new visuals when the app reruns.
    """
    try:
        parts = urlparse(src_url)
        q = parse_qs(parts.query)
        q["cb"] = [str(uuid.uuid4())]  # unique value
        new_q = urlencode(q, doseq=True)
        return urlunparse((parts.scheme, parts.netloc, parts.path, parts.params, new_q, parts.fragment))
    except Exception:
        # Fallback to a simple suffix
        suffix = "&" if "?" in src_url else "?"
        return f"{src_url}{suffix}cb={uuid.uuid4()}"


def main():
    initialize_session_state()

    # Initialize system design session state
    if 'system_design_generated' not in st.session_state:
        st.session_state.system_design_generated = False
    if 'system_design_artifacts' not in st.session_state:
        st.session_state.system_design_artifacts = None

    # Header
    st.markdown("""
    <div class="main-header">
        <h1>📋 AI-Powered Business Analysis Suite</h1>
        <p>Generate BRDs and System Design Artifacts using AI</p>
    </div>
    """, unsafe_allow_html=True)

    # Sidebar configuration
    with st.sidebar:
        st.header("⚙️ Configuration")

        # Environment validation
        missing_vars = validate_environment()
        if missing_vars:
            st.error(f"Missing environment variables: {', '.join(missing_vars)}")
            st.info("Please set up your environment variables in a .env file or system environment.")
            st.stop()
        else:
            st.success("Environment configured ✅")

        # API Configuration
        st.subheader("API Settings")
        api_type = st.selectbox(
            "API Provider",
            options=["openai", "ollama"],
            index=0,
            help="Select your AI API provider"
        )

        if api_type == "openai":
            model = st.selectbox(
                "OpenAI Model",
                options=["gpt-4o-mini", "gpt-4o", "gpt-3.5-turbo"],
                index=0,
                help="Select the OpenAI model to use"
            )
        else:
            model = st.text_input(
                "Ollama Model",
                value="llama3",
                help="Enter the Ollama model name"
            )

        temperature = st.slider(
            "AI Temperature",
            min_value=0.0,
            max_value=1.0,
            value=0.2,
            step=0.1,
            help="Controls randomness in AI responses"
        )

        # Update environment variables
        os.environ["API_TYPE"] = api_type
        if api_type == "openai":
            os.environ["OPENAI_MODEL"] = model
        else:
            os.environ["OLLAMA_MODEL"] = model
        os.environ["AI_TEMPERATURE"] = str(temperature)
        

    # Main navigation tabs
    # NEW: add a third tab for Power BI dashboard
    tab1, tab2, tab3, tab4 = st.tabs(["📝 BRD Generator", "🏗️ System Design Generator", "🛠️ Dashboard Dev Doc", "📈 Power BI Dashboard"])
    with tab1:
        # Main content
        col1, col2 = st.columns([2, 1])

        with col1:
            # Template upload section
            st.header("📋 BRD Template (Optional)")
            template_file = st.file_uploader(
                "Upload BRD Template",
                type=['docx', 'txt', 'md', 'json'],
                help="Upload a custom BRD template with your preferred sections. Supports DOCX (with headings), TXT/MD (with headers), or JSON format.",
                key="template_upload"
            )

            if template_file:
                st.success(f"Template uploaded: {template_file.name}")

                # Process template
                temp_template_path = save_uploaded_file(template_file)
                if temp_template_path:
                    sections = extract_template_sections(temp_template_path)
                    if sections:
                        st.session_state.template_sections = sections
                        st.session_state.custom_template = temp_template_path

                        with st.expander("📝 Template Sections Detected"):
                            for i, section in enumerate(sections, 1):
                                st.write(f"{i}. {section}")

                        st.info(f"✅ Template processed: {len(sections)} sections detected")
                    else:
                        st.warning("Could not extract sections from template. Using default template.")
                        st.session_state.template_sections = None
            else:
                # Show default template
                if st.session_state.template_sections is None:
                    default_sections = [
                        "Executive Summary", "Project Overview", "Scope (In / Out)",
                        "Business Objectives", "Functional Requirements", "Non-Functional Requirements",
                        "Acceptance Criteria", "Assumptions & Constraints", "Risks", "Stakeholders",
                        "Data Flow & Integration", "Milestones & Timeline", "Change Control", "Glossary", "Appendices"
                    ]

                    with st.expander("📝 Default Template Sections"):
                        for i, section in enumerate(default_sections, 1):
                            st.write(f"{i}. {section}")

            st.header("📤 Upload Your BRS Document")

            # File uploader
            uploaded_file = st.file_uploader(
                "Choose a file",
                type=['txt', 'md', 'docx', 'csv', 'xlsx', 'xls', 'json', 'pdf'],
                help="Supported formats: TXT, MD, DOCX, CSV, XLSX, XLS, JSON, PDF"
            )

            if uploaded_file:
                st.success(f"File uploaded: {uploaded_file.name} ({uploaded_file.size} bytes)")

                # File preview
                with st.expander("📖 Preview File Content"):
                    try:
                        temp_path = save_uploaded_file(uploaded_file)
                        if temp_path:
                            preview_text = FileProcessor.extract_text(temp_path)
                            st.text_area(
                                "File Content Preview",
                                value=preview_text[:1000] + ("..." if len(preview_text) > 1000 else ""),
                                height=200,
                                disabled=True
                            )
                            os.unlink(temp_path)  # Clean up temp file
                    except Exception as e:
                        st.error(f"Error previewing file: {str(e)}")

                # Generate BRD button
                if st.button("🚀 Generate BRD", type="primary", use_container_width=True):
                    try:
                        # Save uploaded file
                        temp_file_path = save_uploaded_file(uploaded_file)
                        if not temp_file_path:
                            st.error("Failed to save uploaded file")
                            st.stop()

                        # Progress tracking
                        progress_bar = st.progress(0)
                        status_text = st.empty()

                        def update_progress(percentage, message):
                            progress_bar.progress(percentage)
                            status_text.text(message)

                        # Generate BRD
                        with st.spinner("Generating BRD..."):
                            brd_file_path, brd_content = generate_brd_from_file(
                                temp_file_path,
                                progress_callback=update_progress,
                                template_sections=st.session_state.template_sections
                            )
                            st.success("Draft BRD generated ✅")
                            # --- Review & Compress Step ---
                        with st.spinner("Reviewing and compressing BRD..."):
                            gen = TemplateBRDGenerator(AIProcessor(), st.session_state.template_sections)
                            project_title = "AI-Powered BRD"  # or extract from facts / user input
                            reviewed_brd = gen.review_and_compress(brd_content, project_title)
                            # Second pass for further deduplication
                            reviewed_brd = gen.review_and_compress({"Reviewed BRD": reviewed_brd}, project_title)  # Wrap as dict
                            # 🔑 Save reviewed BRD and point the session state to it
                            reviewed_filename = f"BRD_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                            gen.save_brd_docx({"Reviewed BRD": reviewed_brd}, reviewed_filename)
                            st.session_state.brd_filename = reviewed_filename   # <- makes your existing button download the reviewed file

                        # Store reviewed version in session state
                        st.session_state.reviewed_brd = reviewed_brd
                        

                        # Store in session state
                        st.session_state.brd_generated = True
                        st.session_state.brd_content = brd_content
                        st.session_state.brd_filename = brd_file_path

                        # Clean up temp input file
                        os.unlink(temp_file_path)

                        st.success("🎉 BRD generated successfully!")

                    except Exception as e:
                        st.error(f"Error: {str(e)}")
                        st.info("Please check your file format and API configuration.")

        with col2:
            st.header("📊 Generation Status")

            if st.session_state.brd_generated:
                st.success("✅ BRD Generated")

                # Download button
                if st.session_state.brd_filename and os.path.exists(st.session_state.brd_filename):
                    with open(st.session_state.brd_filename, "rb") as file:
                        btn = st.download_button(
                            label="📥 Download BRD",
                            data=file.read(),
                            file_name=f"BRD_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary",
                            use_container_width=True
                        )

                    # BRD content preview
                    if st.session_state.brd_content:
                        with st.expander("📋 BRD Content Preview"):
                            for section, content in st.session_state.brd_content.items():
                                st.subheader(section)
                                st.write(content[:200] + ("..." if len(content) > 200 else ""))
                                st.divider()

                # Reset button
                if st.button("🔄 Generate New BRD", use_container_width=True):
                    st.session_state.brd_generated = False
                    st.session_state.brd_content = None
                    if st.session_state.brd_filename and os.path.exists(st.session_state.brd_filename):
                        os.unlink(st.session_state.brd_filename)
                    st.session_state.brd_filename = None
                    # Clean up template files
                    if st.session_state.custom_template and os.path.exists(st.session_state.custom_template):
                        os.unlink(st.session_state.custom_template)
                    st.session_state.custom_template = None
                    st.session_state.template_sections = None
                    st.rerun()
            else:
                st.info("Upload a file and click 'Generate BRD' to start")

        # Information section
        with st.expander("ℹ️ How it works"):
            st.markdown("""
            ### BRD Generation Process

            1. **Template Upload (Optional)**: Upload your custom BRD template to define sections
            2. **File Upload**: Upload your Business Requirements Specification (BRS) document
            3. **Text Extraction**: The system extracts text content from your file
            4. **Template Processing**: Custom template sections are extracted and used
            5. **Fact Extraction**: Key information like projects, dates, stakeholders, and requirements are identified
            6. **AI Processing**: Using advanced AI, each BRD section is generated based on extracted facts and template structure
            7. **Document Generation**: A comprehensive BRD is created in DOCX format

            ### Template Formats Supported
            - **DOCX**: Uses heading styles (Heading 1, Heading 2, etc.) to identify sections
            - **TXT/MD**: Recognizes markdown headers (#, ##) or numbered sections (1., 2., etc.)
            - **JSON**: Sections as keys in object or as array of section names

            ### Example Template Formats

            **DOCX Template**: Use Word heading styles for each section

            **Markdown Template**:
            ```
            # Executive Summary
            # Project Overview  
            # Business Requirements
            # Technical Specifications
            ```

            **JSON Template**:
            ```json
            {
              "sections": [
                "Executive Summary",
                "Project Overview", 
                "Business Requirements",
                "Technical Specifications"
              ]
            }
            ```

            ### Supported BRS File Formats
            - **Text Files**: .txt, .md
            - **Documents**: .docx
            - **Spreadsheets**: .csv, .xlsx, .xls
            - **Data**: .json
            - **PDFs**: .pdf (requires pdfminer)

            ### Default BRD Sections (if no template provided)
            - Executive Summary
            - Project Overview
            - Scope (In/Out)
            - Business Objectives
            - Functional Requirements
            - Non-Functional Requirements
            - Acceptance Criteria
            - Assumptions & Constraints
            - Risks
            - Stakeholders
            - Data Flow & Integration
            - Milestones & Timeline
            - Change Control
            - Glossary
            - Appendices
            """)

        # Template creation helper
        with st.expander("🛠️ Template Creation Helper"):
            st.markdown("""
            ### Create Your Own Template

            **Need help creating a custom template?** Here are some common BRD section variations:
            """)

            template_type = st.selectbox(
                "Choose a template type to download:",
                ["Custom", "Agile/Scrum", "Waterfall", "Technical", "Business Process"]
            )

            template_sections_map = {
                "Agile/Scrum": [
                    "Executive Summary", "Product Vision", "User Stories", "Sprint Goals",
                    "Acceptance Criteria", "Definition of Done", "Stakeholders",
                    "Constraints", "Assumptions", "Risks"
                ],
                "Waterfall": [
                    "Project Charter", "Business Case", "Scope Statement", "Requirements",
                    "System Architecture", "Implementation Plan", "Testing Strategy",
                    "Risk Management", "Change Control", "Sign-off"
                ],
                "Technical": [
                    "System Overview", "Technical Requirements", "Architecture Design",
                    "API Specifications", "Database Schema", "Security Requirements",
                    "Performance Criteria", "Integration Points", "Deployment Plan"
                ],
                "Business Process": [
                    "Process Overview", "Current State", "Future State", "Gap Analysis",
                    "Business Rules", "Process Flow", "Roles & Responsibilities",
                    "Performance Metrics", "Implementation Roadmap"
                ]
            }

            if template_type != "Custom" and template_type in template_sections_map:
                sections = template_sections_map[template_type]

                # Generate downloadable template
                template_content = "\n".join([f"# {section}" for section in sections])

                st.download_button(
                    label=f"📥 Download {template_type} Template",
                    data=template_content,
                    file_name=f"{template_type.lower().replace('/', '_')}_template.md",
                    mime="text/markdown"
                )

                st.markdown("**Template Preview:**")
                for i, section in enumerate(sections, 1):
                    st.write(f"{i}. {section}")

    with tab2:
        system_design_interface()
        # Display artifacts and download buttons
        if st.session_state.get('system_design_generated', False):
            display_system_design_results()
# NEW: Update tab definition to include the new tab
    with tab3:
        dashboard_dev_doc_interface()
    # -------------------------------
    # NEW TAB: Power BI Dashboard
    # -------------------------------
    with tab4:
        st.header("📈 Power BI Dashboard")

        st.markdown(
            "This tab embeds your Power BI report. "
            "Use **Auto-Refresh** to periodically reload the iframe so it reflects **incremental refresh** "
            "and any dataset/report changes once they’re available in the Power BI Service."
        )

        # Controls
        with st.expander("🔧 Display & Refresh Settings", expanded=True):
            colA, colB, colC, colD = st.columns(4)

            with colA:
                st.session_state.pbi_autorefresh_enabled = st.checkbox(
                    "Enable Auto-Refresh",
                    value=st.session_state.pbi_autorefresh_enabled,
                    help="Reruns the app at the interval below to reload the embedded report."
                )
                st.session_state.pbi_cache_bust = st.checkbox(
                    "Force Cache-Bust on reload",
                    value=st.session_state.pbi_cache_bust,
                    help="Appends a unique parameter to the URL each time so the iframe fully reloads."
                )

            with colB:
                st.session_state.pbi_autorefresh_seconds = st.number_input(
                    "Auto-Refresh Interval (seconds)",
                    min_value=30,
                    max_value=3600,
                    value=st.session_state.pbi_autorefresh_seconds,
                    step=30
                )

            with colC:
                st.session_state.pbi_height_px = st.number_input(
                    "Viewport Height (px)",
                    min_value=400,
                    max_value=2000,
                    value=st.session_state.pbi_height_px,
                    step=50
                )
            with colD:
                st.session_state.pbi_width_px = st.number_input(
                    "Viewport Width (px)",
                    min_value=400,
                    max_value=2000,
                    value=st.session_state.pbi_width_px,
                    step=50
                )

            # Manual refresh button
            refresh_now = st.button("🔄 Refresh Now")

        # Your shared Power BI public link (iframe src)
        base_powerbi_src = (
            "https://app.powerbi.com/reportEmbed?reportId=e6eb8cfa-62e8-47e0-926f-316eb03db390&autoAuth=true&ctid=e55e9f72-ac55-436c-aa55-a40f0dc07184&actionBarEnabled=true"        )

        # Handle auto-refresh by triggering reruns
        if st.session_state.pbi_autorefresh_enabled:
            # st_autorefresh returns a counter; we don't need it, just call to enable reruns
            st.experimental_rerun  # hint for readability; execution below:
            st.experimental_set_query_params()  # no-op; keeps URL tidy
            st.autorefresh = st.experimental_singleton(lambda: True)  # dummy to avoid lints
            st_autorefresh = st.experimental_memo(lambda: True)       # dummy to avoid lints
            # actual API:
            st.experimental_rerun  # ensure imported symbol exists
            st_autorefresh_counter = st.experimental_get_query_params()  # harmless
            st_autorefresh = st.experimental_rerun  # no effect
            # Correct method:
            _ = st.experimental_data_editor if False else None  # silence warnings
            st_autorefresh_obj = st.experimental_show if False else None  # silence warnings
            # Use the built-in st_autorefresh:
            _ = st.autorefresh if hasattr(st, "autorefresh") else None
        # Streamlit added st_autorefresh utility; call when available
        try:
            st_autorefresh = getattr(st, "autorefresh")
            if st.session_state.pbi_autorefresh_enabled and callable(st_autorefresh):
                st_autorefresh(interval=st.session_state.pbi_autorefresh_seconds * 1000, key="pbi_autorefresh_key")
        except Exception:
            pass  # Older Streamlit versions may not have st.autorefresh; the manual refresh button still works.

        # Build the src with optional cache-busting
        src = base_powerbi_src
        if st.session_state.pbi_cache_bust or refresh_now:
            src = _cache_busted_url(src)

        # Render iframe (responsive container)
        # We set a fixed height via style so users can adjust; width is 90% responsive.
        html_iframe = f"""
            <div style="position: relative; width: {st.session_state.pbi_width_px}px; height: {st.session_state.pbi_height_px}px;">
        <iframe title="AI PowerBI dashboard Public"
                width="{st.session_state.pbi_width_px}"
                height="{st.session_state.pbi_height_px}"
                src="{src}"
                frameborder="0"
                allowFullScreen="true"></iframe>
    </div>
        """
        components.html(html_iframe, height=st.session_state.pbi_height_px + 10, width=st.session_state.pbi_width_px + 10, scrolling=True)

        st.info(
            "💡 Notes:\n"
            "- The embedded report reflects **incremental refresh** and any dataset updates once published in the Power BI Service.\n"
            "- **Auto-Refresh** here simply reloads the iframe at the chosen interval so changes appear without manual reload.\n"
            "- Use **Refresh Now** anytime to immediately force a cache-busted reload."
        )
# NEW: Dashboard Developer Document Generator Class
class CopilotDashboardGenerator:
    """Generates a concise Dashboard Developer Document optimized for Microsoft Copilot."""

    def __init__(self, ai: AIProcessor):
        self.ai = ai

    def generate_copilot_dashboard_doc(self, brd_content: Dict[str, str]) -> Dict[str, str]:
        """Generate a concise, Copilot-friendly dashboard document."""
        
        # Convert BRD to text
        brd_text = "\n\n".join(f"{sec}: {content}" for sec, content in brd_content.items())
        
        # Define focused sections for Copilot
        sections = {
            "Dashboard Overview": self._generate_overview,
            "Data Source Configuration": self._generate_data_config,
            "Page 1 - Executive Summary": self._generate_page1,
            "Page 2 - Trend Analysis": self._generate_page2,
            "Page 3 - Performance Metrics": self._generate_page3,
            "Page 4 - Detailed Analytics": self._generate_page4,
            "Global Filters": self._generate_filters,
            "Formatting Guidelines": self._generate_formatting
        }
        
        doc = {}
        for section_name, generator_func in sections.items():
            doc[section_name] = generator_func(brd_text)
        
        return doc

    def _generate_overview(self, brd_text: str) -> str:
        """Generate concise overview section."""
        system_prompt = """You are creating instructions for Microsoft Copilot to build a Power BI dashboard.
Be extremely concise and specific. Use bullet points and clear directives."""
        
        user_prompt = f"""Based on this BRD, create a 3-4 sentence dashboard overview:
{brd_text[:1500]}

Format:
- Dashboard purpose (1 sentence)
- Key metrics to track (1 sentence)
- Target users (1 sentence)
- Update frequency (1 sentence)"""
        
        return self.ai.call_ai(system_prompt, user_prompt)

    def _generate_data_config(self, brd_text: str) -> str:
        """Generate data source configuration."""
        return """Data Source: Connect to the provided CSV file
Required Columns:
• EventID, MachineID, EventTimestamp
• EventType, AlarmCode, AlarmDescription
• Severity, OperatorID, Status
• ResolutionTime, DurationSeconds
• TemperatureReading, VibrationLevel
• ErrorFlag, SparePartUsed

Create Calculated Columns:
1. Shift = IF(HOUR([EventTimestamp]) >= 5 AND HOUR([EventTimestamp]) < 17, "Day Shift", "Night Shift")
2. ResolutionMinutes = [DurationSeconds] / 60
3. EventDate = DATE([EventTimestamp])
4. EventHour = HOUR([EventTimestamp])"""

    def _generate_page1(self, brd_text: str) -> str:
        """Generate Page 1 specifications."""
        return """Title: Executive Summary

Layout (Top to Bottom):
1. KPI Cards Row (4 cards):
   • Card 1: Total Events - COUNT(EventID)
   • Card 2: Critical Alarms (24hr) - COUNT(EventID) WHERE Severity="Critical" AND EventTimestamp >= NOW()-1
   • Card 3: Avg Resolution Time - AVERAGE(ResolutionMinutes) FORMAT: "0.0 mins"
   • Card 4: Open Events % - COUNT(Status="Open")/COUNT(EventID) FORMAT: Percentage

2. Line Chart - Events Over Time:
   • X-axis: EventTimestamp (Date hierarchy)
   • Y-axis: Count of EventID
   • Legend: Severity (Critical=Red, High=Orange, Low=Green)
   • Title: "Event Trend by Severity"

3. Donut Chart - Event Type Distribution:
   • Values: Count of EventID
   • Legend: EventType
   • Show data labels as percentages"""

    def _generate_page2(self, brd_text: str) -> str:
        """Generate Page 2 specifications."""
        return """Title: Trend Analysis

Layout:
1. Stacked Column Chart - Shift Performance:
   • X-axis: Shift
   • Y-axis: Count of EventID
   • Legend: Severity
   • Title: "Events by Shift and Severity"

2. Area Chart - Hourly Distribution:
   • X-axis: EventHour (0-23)
   • Y-axis: Count of EventID
   • Series: EventType
   • Title: "24-Hour Event Pattern"

3. Matrix Visual:
   • Rows: MachineID
   • Columns: Shift
   • Values: Count of EventID (heat map formatting)
   • Conditional formatting: Red for >10 events"""

    def _generate_page3(self, brd_text: str) -> str:
        """Generate Page 3 specifications."""
        return """Title: Performance Metrics

Layout:
1. Scatter Chart - Machine Health:
   • X-axis: TemperatureReading
   • Y-axis: VibrationLevel
   • Size: DurationSeconds
   • Color: ErrorFlag (True=Red, False=Green)
   • Title: "Temperature vs Vibration Analysis"

2. Horizontal Bar Chart - Machine Downtime:
   • Y-axis: MachineID (Top 10)
   • X-axis: SUM(DurationSeconds)/3600 (hours)
   • Sort: Descending by value
   • Title: "Top 10 Machines by Downtime"

3. Table - Machine Statistics:
   • Columns: MachineID | Total Events | Critical Events | Avg Resolution Time
   • Sort by: Total Events (descending)
   • Limit: Top 20 machines"""

    def _generate_page4(self, brd_text: str) -> str:
        """Generate Page 4 specifications."""
        return """Title: Detailed Analytics

Layout:
1. Clustered Bar Chart - Operator Performance:
   • Y-axis: OperatorID
   • X-axis: Average ResolutionMinutes
   • Sort: Ascending by resolution time
   • Title: "Operator Efficiency Ranking"

2. Table - Operator Details:
   • Columns: OperatorID | Events Handled | Avg Resolution (mins) | Spare Parts Used
   • Sort by: Events Handled (descending)

3. KPI Cards Row:
   • Best Performer: MIN(Average ResolutionMinutes) with OperatorID
   • Most Active: MAX(COUNT(EventID)) with OperatorID
   • Parts Usage Rate: COUNT(SparePartUsed NOT NULL)/COUNT(EventID)"""

    def _generate_filters(self, brd_text: str) -> str:
        """Generate global filters specification."""
        return """Global Filters (Apply to All Pages):

1. Date Range Slicer:
   • Field: EventTimestamp
   • Type: Between
   • Default: Last 30 days

2. Machine Multi-Select:
   • Field: MachineID
   • Type: Dropdown (multi-select enabled)

3. Severity Buttons:
   • Field: Severity
   • Type: Tile slicer
   • Colors: Critical=Red, High=Orange, Low=Green

4. Event Type Dropdown:
   • Field: EventType
   • Type: Single select dropdown

5. Status Toggle:
   • Field: Status
   • Type: Button slicer
   • Default: All selected"""

    def _generate_formatting(self, brd_text: str) -> str:
        """Generate formatting guidelines."""
        return """Formatting Standards:

Colors:
• Critical: #FF0000 (Red)
• High: #FFA500 (Orange)  
• Low: #008000 (Green)
• Background: #F8F9FA (Light Gray)
• Primary: #0078D4 (Blue)

Typography:
• Headers: Segoe UI, 16pt, Bold
• Body: Segoe UI, 12pt
• KPI Values: Segoe UI, 24pt, Bold

Conditional Formatting Rules:
• Resolution Time: >60 mins = Red, 30-60 = Orange, <30 = Green
• Event Count: >100 = Red, 50-100 = Orange, <50 = Green
• Apply to all tables and matrices

General:
• Enable tooltips on all visuals
• Use thousands separators
• Time format: HH:MM
• Date format: MM/DD/YYYY"""

    def _clean_content(self, text: str) -> str:
        """Clean and format text content."""
        # Remove markdown formatting
        text = re.sub(r"\*\*\*(.*?)\*\*\*", r"\1", text)
        text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
        text = re.sub(r"\*(.*?)\*", r"\1", text)
        text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def save_docx(self, doc: Dict[str, str], output_path: str):
        """Save document as DOCX file."""
        word_doc = DocxDocument()
        
        # Title
        title = word_doc.add_heading("Power BI Dashboard Development Guide for Microsoft Copilot", 0)
        title.alignment = 1
        word_doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        word_doc.add_paragraph("Instructions: Use this document with Microsoft Copilot to generate Power BI dashboard")
        
        word_doc.add_page_break()
        
        # Add sections
        for section, content in doc.items():
            word_doc.add_heading(section, level=1)
            # Process content by paragraphs
            for para in content.split("\n"):
                if para.strip():
                    if para.startswith("•"):
                        # Add as bullet point
                        word_doc.add_paragraph(para, style='List Bullet')
                    else:
                        word_doc.add_paragraph(para)
        
        word_doc.save(output_path)


def dashboard_dev_doc_interface():
    """Streamlit interface for dashboard document generation."""
    st.header("📊 Copilot-Ready Dashboard Document Generator")
    st.markdown("Generate a concise document optimized for Microsoft Copilot to create Power BI dashboards.")
    
    # Input method for BRD
    st.subheader("📄 Input Business Requirements")
    brd_input_method = st.radio(
        "Choose BRD input method:",
        ["Use Generated BRD", "Paste BRD Text", "Upload BRD File"],
        horizontal=True
    )
    
    brd_content = None
    
    if brd_input_method == "Use Generated BRD":
        if st.session_state.get('brd_generated') and st.session_state.get('brd_content'):
            brd_content = st.session_state.brd_content
            with st.expander("📖 Current BRD Content"):
                brd_text = "\n\n".join(f"{sec}: {content}" for sec, content in brd_content.items())
                st.text_area("BRD Preview", value=brd_text[:1000] + "..." if len(brd_text) > 1000 else brd_text, 
                           height=150, disabled=True)
        else:
            st.warning("⚠️ No BRD has been generated. Please generate a BRD first or choose another input method.")
            return
    
    elif brd_input_method == "Paste BRD Text":
        brd_text = st.text_area(
            "Paste your BRD content here:",
            height=300,
            placeholder="Paste your business requirements document here..."
        )
        if brd_text.strip():
            brd_content = {"Business Requirements": brd_text}
        else:
            st.info("💡 Paste your BRD content to proceed")
            return
    
    elif brd_input_method == "Upload BRD File":
        uploaded_brd = st.file_uploader(
            "Upload BRD File",
            type=['txt', 'md', 'docx', 'pdf', 'json'],
            help="Supported formats: TXT, Markdown, Word, PDF, JSON"
        )
        if uploaded_brd:
            try:
                brd_path = save_uploaded_file(uploaded_brd)
                if brd_path:
                    brd_text = FileProcessor.extract_text(brd_path)
                    os.unlink(brd_path)
                    brd_content = {"Uploaded BRD": brd_text}
                    st.success("✅ File uploaded successfully!")
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")
                return
        else:
            st.info("📁 Upload a BRD file to proceed")
            return
    
    # Configuration options
    with st.expander("⚙️ Dashboard Configuration", expanded=False):
        col1, col2 = st.columns(2)
        with col1:
            num_pages = st.number_input("Number of Dashboard Pages", min_value=1, max_value=10, value=4)
            include_dax = st.checkbox("Include DAX Formulas", value=True)
        with col2:
            theme = st.selectbox("Visual Theme", ["Professional", "Modern", "Minimal", "Dark"])
            include_bookmarks = st.checkbox("Include Bookmarks", value=False)
    
    # Generate button
    if st.button("🚀 Generate Copilot-Ready Document", use_container_width=True, type="primary"):
        with st.spinner("Generating optimized dashboard document..."):
            try:
                # Initialize generators
                ai = AIProcessor()
                generator = CopilotDashboardGenerator(ai)
                
                # Generate document
                doc_content = generator.generate_copilot_dashboard_doc(brd_content)
                
                # Save to file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                    generator.save_docx(doc_content, tmp_file.name)
                    st.session_state.dashboard_dev_filename = tmp_file.name
                    st.session_state.dashboard_dev_content = doc_content
                    st.session_state.dashboard_dev_generated = True
                
                st.success("✅ Dashboard document generated successfully!")
                st.balloons()
                
            except Exception as e:
                st.error(f"❌ Error generating document: {str(e)}")
                return
    
    # Display results
    if st.session_state.get('dashboard_dev_generated'):
        st.divider()
        
        # Preview sections
        st.subheader("📋 Document Preview")
        
        # Create tabs for each section
        tabs = st.tabs(list(st.session_state.dashboard_dev_content.keys()))
        
        for tab, (section, content) in zip(tabs, st.session_state.dashboard_dev_content.items()):
            with tab:
                st.text(content)
        
        # Download section
        st.divider()
        col1, col2, col3 = st.columns([2, 2, 1])
        
        with col1:
            with open(st.session_state.dashboard_dev_filename, "rb") as f:
                st.download_button(
                    label="📥 Download for Copilot (.DOCX)",
                    data=f,
                    file_name=f"Copilot_Dashboard_Guide_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        
        with col2:
            # Option to copy to clipboard (for web-based Copilot)
            if st.button("📋 Copy All Content", use_container_width=True):
                all_content = "\n\n".join([f"## {sec}\n{content}" for sec, content in st.session_state.dashboard_dev_content.items()])
                st.code(all_content, language="markdown")
                st.info("📌 Content displayed above - select all and copy to use with Copilot")
        
        with col3:
            if st.button("🔄 New Doc", use_container_width=True):
                # Reset state
                st.session_state.dashboard_dev_generated = False
                st.session_state.dashboard_dev_content = None
                if st.session_state.get('dashboard_dev_filename') and os.path.exists(st.session_state.dashboard_dev_filename):
                    os.unlink(st.session_state.dashboard_dev_filename)
                st.session_state.dashboard_dev_filename = None
                st.rerun()
        
        # Usage instructions
        with st.expander("📘 How to use with Microsoft Copilot", expanded=False):
            st.markdown("""
            ### Using this document with Microsoft Copilot:
            
            1. **Open Power BI Desktop** with Copilot enabled
            2. **Load your data source** (CSV file mentioned in the document)
            3. **Open Copilot pane** in Power BI
            4. **Copy each page section** from this document
            5. **Paste into Copilot** with instruction: "Create this dashboard page"
            6. **Review and adjust** the generated visuals
            
            ### Tips for best results:
            - Process one page at a time
            - Ensure data is loaded before using Copilot
            - Verify calculated columns are created first
            - Apply global filters after all pages are created
            """)

# Initialize session state variables
if 'dashboard_dev_generated' not in st.session_state:
    st.session_state.dashboard_dev_generated = False
    st.session_state.dashboard_dev_content = None
    st.session_state.dashboard_dev_filename = None



if __name__ == "__main__":
    main()
